"""Extract text from PDF / DOCX / XLSX / DOC / ZIP. Spec §6.9.

Returns ExtractedText. Never raises — failures become extraction_status values.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import cast

from collector.audit_log import get_logger
from collector.models import ExtractedText, ExtractionStatus

log = get_logger(__name__)

_SCANNED_TEXT_THRESHOLD = 50  # chars
_TELUGU_RANGE = (0x0C00, 0x0C7F)
_TELUGU_RATIO_THRESHOLD = 0.05
_LATIN_RATIO_THRESHOLD = 0.6


def _guess_language(text: str) -> str | None:
    if not text.strip():
        return None
    total = len(text)
    if total == 0:
        return None
    telugu = sum(1 for ch in text if _TELUGU_RANGE[0] <= ord(ch) <= _TELUGU_RANGE[1])
    latin = sum(1 for ch in text if 32 <= ord(ch) <= 126)
    if telugu / total > _TELUGU_RATIO_THRESHOLD:
        return "te"
    if latin / total > _LATIN_RATIO_THRESHOLD:
        return "en"
    return "mixed"


def _write_text(out_dir: Path, document_id: str, text: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{document_id}.txt"
    out_path.write_text(text, encoding="utf-8")
    return out_path


def _extract_pdf(file_path: Path) -> tuple[str, int | None, str | None]:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        return "", None, f"pymupdf_not_installed:{exc!s}"
    try:
        with fitz.open(file_path) as doc:
            if doc.is_encrypted and not doc.authenticate(""):
                return "", doc.page_count, "encrypted"
            chunks = [page.get_text("text") for page in doc]
            return "\n".join(chunks), doc.page_count, None
    except Exception as exc:
        return "", None, f"pdf_error:{exc!s}"


def _extract_docx(file_path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        return "", f"python-docx_not_installed:{exc!s}"
    try:
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(paragraphs), None
    except Exception as exc:
        return "", f"docx_error:{exc!s}"


def _extract_xlsx(file_path: Path) -> tuple[str, str | None]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        return "", f"openpyxl_not_installed:{exc!s}"
    try:
        wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
        chunks: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            chunks.append(f"# Sheet: {sheet_name}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 100:
                    chunks.append("...")
                    break
                chunks.append(" | ".join("" if v is None else str(v) for v in row))
        return "\n".join(chunks), None
    except Exception as exc:
        return "", f"xlsx_error:{exc!s}"


def _safe_zip_extract(
    file_path: Path,
    workdir: Path,
    allowed_exts: tuple[str, ...],
    max_files: int = 200,
    max_total_bytes: int = 500 * 1024 * 1024,
) -> list[Path]:
    out: list[Path] = []
    workdir.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(file_path) as zf:
            members = zf.infolist()
            if len(members) > max_files:
                log.warning("zip_too_many_members", path=str(file_path), n=len(members))
                return out
            total = 0
            for member in members:
                name = member.filename
                if name.startswith("/") or ".." in Path(name).parts:
                    continue
                ext = Path(name).suffix.lstrip(".").lower()
                if ext not in allowed_exts:
                    continue
                if member.file_size > max_total_bytes - total:
                    log.warning("zip_size_cap_reached", path=str(file_path))
                    break
                target = workdir / Path(name).name
                with zf.open(member) as src, target.open("wb") as dst:
                    while True:
                        chunk = src.read(64 * 1024)
                        if not chunk:
                            break
                        dst.write(chunk)
                        total += len(chunk)
                        if total > max_total_bytes:
                            target.unlink(missing_ok=True)
                            log.warning("zip_total_size_exceeded", path=str(file_path))
                            return out
                out.append(target)
    except zipfile.BadZipFile as exc:
        log.warning("bad_zip", path=str(file_path), error=str(exc))
    return out


def extract(
    *,
    document_id: str,
    file_path: Path,
    file_extension: str,
    processed_root: Path,
    allowed_exts: tuple[str, ...] = ("pdf", "doc", "docx", "xls", "xlsx"),
) -> ExtractedText:
    out_dir = processed_root
    text = ""
    page_count: int | None = None
    status: str = "ok"
    err: str | None = None
    manual = False

    ext = file_extension.lower()
    if ext == "pdf":
        text, page_count, err = _extract_pdf(file_path)
        if err == "encrypted":
            status, manual = "encrypted", True
        elif err:
            status = "failed"
        elif page_count and len(text.strip()) < _SCANNED_TEXT_THRESHOLD:
            status, manual = "scanned_or_no_text", True
    elif ext == "docx":
        text, err = _extract_docx(file_path)
        if err:
            status = "failed"
    elif ext in ("xlsx", "xls"):
        text, err = _extract_xlsx(file_path)
        if err:
            status = "failed"
    elif ext == "doc":
        status, manual, err = "unsupported", True, "legacy_doc_unsupported_in_v1"
    elif ext == "zip":
        sub_files = _safe_zip_extract(file_path, out_dir / f"_zip_{document_id}", allowed_exts)
        sub_texts: list[str] = []
        for sf in sub_files:
            sub_ext = sf.suffix.lstrip(".").lower()
            sub_text = ""
            if sub_ext == "pdf":
                sub_text, _, _ = _extract_pdf(sf)
            elif sub_ext == "docx":
                sub_text, _ = _extract_docx(sf)
            elif sub_ext in ("xlsx", "xls"):
                sub_text, _ = _extract_xlsx(sf)
            if sub_text:
                sub_texts.append(f"# {sf.name}\n{sub_text}")
        text = "\n\n".join(sub_texts)
        if not text:
            status, manual = "scanned_or_no_text", True
    else:
        status, manual, err = "unsupported", True, f"unsupported_extension:{ext}"

    extracted_path = _write_text(out_dir, document_id, text)
    preview = text[:1000]
    lang = _guess_language(text) if text else None

    return ExtractedText(
        document_id=document_id,
        extracted_text_path=extracted_path,
        text_preview=preview,
        page_count=page_count,
        language_guess=lang,
        extraction_status=cast(ExtractionStatus, status),
        extraction_error=err,
        manual_review_required=manual,
    )


_REPEATED_WS_RX = re.compile(r"[ \t]+")


def normalize_whitespace(text: str) -> str:
    return _REPEATED_WS_RX.sub(" ", text).strip()
